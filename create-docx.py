from docx import Document

def create_word_document(data, output_file):
    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Sample Data Report', 0)

    # Add data to the document
    for item in data:
        doc.add_paragraph(item)

    # Save the document to a file
    doc.save(output_file)
    print(f"Word document '{output_file}' created successfully.")

if __name__ == "__main__":
    # Sample data to be included in the Word document
    data = [
        "This is the first paragraph of data.",
        "This is the second paragraph of data.",
        "You can customize this data as needed.",
    ]

    # Specify the output file name
    output_file = "sample_data_report.docx"

    # Generate the Word document
    create_word_document(data, output_file)